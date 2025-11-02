import json
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.sync_api import sync_playwright


def add_header_with_style(doc, text, level=1):
    """Add a styled header to the document."""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(52, 46, 173)
    return heading


def add_table_with_data(doc, headers, rows, col_widths=None):
    """Create a formatted table with headers and data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    # Add data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    # This function adds a hyperlink with proper formatting
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = paragraph._element
    hyperlink_tag = hyperlink.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink')
    hyperlink_tag.set('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id)
    
    run_element = hyperlink_tag.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    run_properties = run_element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
    
    # Add underline and color for hyperlink
    u_element = run_properties.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u')
    u_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
    run_properties.append(u_element)
    
    color_element = run_properties.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color')
    color_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0563C1')
    run_properties.append(color_element)
    
    run_element.append(run_properties)
    
    text_element = run_element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
    text_element.text = text
    run_element.append(text_element)
    
    hyperlink_tag.append(run_element)
    hyperlink.append(hyperlink_tag)


def add_summary_table(doc, data_dict):
    """Create a two-column summary table."""
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (key, value) in enumerate(data_dict.items()):
        row_cells = table.rows[i].cells
        row_cells[0].text = key
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = str(value)
    
    return table


def add_summary_table_with_note(doc, data_dict, note_text=None, note_url=None):
    """Create a two-column summary table with optional note row."""
    table = doc.add_table(rows=len(data_dict), cols=2)
    table.style = 'Light Grid Accent 1'
    
    for i, (key, value) in enumerate(data_dict.items()):
        row_cells = table.rows[i].cells
        row_cells[0].text = key
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].text = str(value)
    
    # Add note row if provided
    if note_text:
        row_cells = table.add_row().cells
        # Merge both cells for the note
        row_cells[0].merge(row_cells[1])
        paragraph = row_cells[0].paragraphs[0]
        
        if note_url:
            # Split note text at the link
            parts = note_text.split('Attendance Details')
            if len(parts) == 2:
                run = paragraph.add_run(parts[0])
                run.font.bold = True
                add_hyperlink(paragraph, note_url, 'Attendance Details')
                if len(parts[1]) > 0:
                    run2 = paragraph.add_run(parts[1])
                    run2.font.bold = True
            else:
                run = paragraph.add_run(note_text)
                run.font.bold = True
        else:
            run = paragraph.add_run(note_text)
            run.font.bold = True
    
    return table


def capture_chart_image(html_path, output_image_path):
    """Capture the Chart.js visualization from HTML as an image using Playwright."""
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch()
            page = browser.new_page()
            
            # Load the HTML file
            page.goto(f'file:///{html_path}')
            
            # Wait for Chart.js to render
            page.wait_for_timeout(3000)
            
            # Locate the chart canvas and take screenshot
            chart_element = page.query_selector('#sprintVelocityChart')
            if chart_element:
                chart_element.screenshot(path=str(output_image_path))
                browser.close()
                return True
            
            browser.close()
            return False
    except Exception as e:
        print(f"Error capturing chart image: {e}")
        return False


def generate_doc_report(json_path, output_path, html_path=None):
    """Generate a .docx report from JSON data matching the HTML format."""
    
    # Load JSON data
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create document
    doc = Document()
    
    # Set A4 page size and narrow margins
    section = doc.sections[0]
    
    # A4 dimensions: 21 cm x 29.7 cm
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    
    # Narrow margins: 0.5 inches (1.27 cm) all around
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Header Section
    title = doc.add_heading('Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(108, 99, 255)
    
    subtitle = doc.add_paragraph('A comprehensive performance overview')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)
    doc.add_paragraph()  # Spacing
    
    # General Information
    add_header_with_style(doc, 'General Information', level=2)
    general_info = {
        'Name': data['employee_name'],
        'Team': data['team'],
        'Evaluation Period': data['evaluation_period']
    }
    add_summary_table(doc, general_info)
    doc.add_paragraph()
    
    # Attendance Summary
    add_header_with_style(doc, 'Attendance Summary', level=2)
    attendance = data['attendance_summary']
    attendance_info = {
        'Total Working Days': attendance['total_days'],
        'Days Present': attendance['present_days'],
        'Days Absent': attendance['absent_days']
    }
    note_text = 'Note: Check out complete attendance details on Sharepoint: Attendance Details'
    note_url = 'https://rihalom598.sharepoint.com/:x:/s/CodelineAffairs/ES6xjksQQ4dNqy9VzbRFa8sBwI9eShNDFoJgmnIeRzDPHA?e=0ns2m5'
    add_summary_table_with_note(doc, attendance_info, note_text, note_url)
    doc.add_paragraph()
    
    # Sprint Velocity
    add_header_with_style(doc, 'Sprint Velocity', level=2)
    
    # Try to capture and embed chart image
    chart_added = False
    if html_path and Path(html_path).exists():
        temp_image_path = Path(output_path).parent / f"temp_chart_{Path(json_path).stem}.png"
        if capture_chart_image(html_path, temp_image_path):
            try:
                # Calculate available width (page width minus margins)
                available_width = section.page_width - section.left_margin - section.right_margin
                doc.add_picture(str(temp_image_path), width=available_width)
                doc.add_paragraph()  # Spacing after image
                chart_added = True
                # Clean up temporary image
                temp_image_path.unlink()
            except Exception as e:
                print(f"Error adding chart image: {e}")
    
    if not chart_added:
        doc.add_paragraph("(Chart visualization requires HTML file)")
        doc.add_paragraph()
    
    # Add chart data as table
    sprint_headers = ['Sprint', 'Committed Work (%)', 'Delivered Work (%)', 'Plagiarism']
    sprint_rows = []
    for sprint in data['sprint_velocity']:
        plagiarism = sprint.get('plagiarism', 'No')
        sprint_rows.append([
            sprint['sprint'],
            sprint['committed'],
            sprint['delivered'],
            plagiarism
        ])
    
    table = add_table_with_data(doc, sprint_headers, sprint_rows)
    
    # Add plagiarism note row
    note_row = table.add_row()
    merged_cell = note_row.cells[0].merge(note_row.cells[3])
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('Note: Some commitments are not considered as delivered if there is a case of plagiarism.')
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True
    
    # Add SharePoint link note row
    link_row = table.add_row()
    merged_cell = link_row.cells[0].merge(link_row.cells[3])
    paragraph = merged_cell.paragraphs[0]
    run = paragraph.add_run('Note: Check out the complete Sprint Details on Sharepoint: ')
    run.font.bold = True
    add_hyperlink(paragraph, 'https://rihalom598.sharepoint.com/:x:/s/CodelineAffairs/EehfSy55bmhGnWc5rCqXoOsB0EczeURsqmlCdKgH55vl6A?e=anrUrL', 'Sprint Details')
    
    doc.add_paragraph()
    
    # Monthly Evaluation Outcomes
    add_header_with_style(doc, 'Monthly Evaluation Outcomes', level=2)
    
    overall = doc.add_paragraph()
    overall.add_run('Overall Performance: ').bold = True
    overall.add_run(data['monthly_evaluation']['Overall Performance']).bold = True
    doc.add_paragraph()
    
    # Monthly Progress
    add_header_with_style(doc, 'Monthly Progress', level=3)
    monthly_headers = ['Month', 'Performance (Out of 100%)', 'Notes']
    monthly_rows = []
    for month_data in data['monthly_evaluation']['Monthly Progress']:
        monthly_rows.append([
            month_data['month'],
            month_data['percentage'],
            month_data['notes']
        ])
    
    add_table_with_data(doc, monthly_headers, monthly_rows)
    doc.add_paragraph()
    
    # Key Strengths
    add_header_with_style(doc, 'Key Strengths', level=3)
    for strength in data['monthly_evaluation']['Key Strengths']:
        doc.add_paragraph(strength, style='List Bullet')
    
    # Areas for Improvement
    add_header_with_style(doc, 'Areas for Improvement', level=3)
    for improvement in data['monthly_evaluation']['Areas for Improvement']:
        doc.add_paragraph(improvement, style='List Bullet')
    doc.add_paragraph()
    
    # Trainer's Feedback
    add_header_with_style(doc, "Trainer's Feedback", level=2)
    for feedback in data['trainers_feedback']:
        doc.add_paragraph(feedback, style='List Bullet')
    
    # Save document
    doc.save(output_path)
    print(f"Document generated successfully: {output_path}")


def main():
    """Main function to generate .docx reports from JSON files."""
    # Resolve paths relative to script location
    script_dir = Path(__file__).resolve().parent.parent
    json_dir = script_dir / 'transformed_data' / 'individual_reports'
    html_dir = script_dir / 'output_reports_html'
    output_dir = script_dir / 'output_reports_doc'
    
    # Create output directory if it doesn't exist
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Process all JSON files
    json_files = list(json_dir.glob('*.json'))
    
    if not json_files:
        print(f"No JSON files found in {json_dir}")
        return
    
    print(f"Found {len(json_files)} JSON file(s) to process")
    
    for json_file in json_files:
        try:
            output_file = output_dir / f"{json_file.stem}.docx"
            html_file = html_dir / f"{json_file.stem}.html"
            
            # Pass HTML path if it exists
            html_path = str(html_file) if html_file.exists() else None
            
            generate_doc_report(str(json_file), str(output_file), html_path)
        except Exception as e:
            print(f"Error processing {json_file.name}: {e}")


if __name__ == "__main__":
    main()
